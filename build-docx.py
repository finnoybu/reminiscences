"""
Build DOCX for "A Sailor's Reminiscences from the Days of the Sailships"

Assembles markdown chapters in order, strips frontmatter, adds chapter
headings with hero illustrations (converted to grayscale JPEG for print),
runs Pandoc, then post-processes the DOCX with custom title-page layout,
typography (Georgia 11 / Calibri headings), 6 x 9 trade paperback page
size, and centered page numbers.

Adapted from finnoybu-fiction's build-docx.py — single book (memoir,
not trilogy) and chapter heroes preserved (memoir is illustrated;
fiction print is text-only).
"""

import os
import re
import subprocess
import yaml
from PIL import Image

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

ROOT = os.path.dirname(os.path.abspath(__file__))
CONTENT_DIR = os.path.join(ROOT, "content", "en")
FRONT_MATTER_PATH = os.path.join(ROOT, "content", "print", "00-front-matter.md")
OUTPUT_DIR = os.path.join(ROOT, "output")
PUBLIC_IMAGES = os.path.join(ROOT, "public", "images")
PRINT_IMAGES_DIR = os.path.join(OUTPUT_DIR, "images-print")
OUTPUT_FILE = "a-sailors-reminiscences.docx"

BOOK_TITLE = "A Sailor's Reminiscences from the Days of the Sailships"
AUTHOR = "Olavus Vullum Bjørnson Vestbø"
YEAR = "2026"
TRANSLATOR_LINE = "Translated from Norwegian by B. O. Berge"
TRANSCRIBER_LINE = "Transcribed by Kenneth Tannenbaum"
PUBLISHER = "Finnoybu Press"
LOGO_PATH = os.path.join(PUBLIC_IMAGES, "finnoybu-press-logo.png")

# Files in public/images that aren't chapter heroes — exclude from print
# image preparation so we don't ship the cover, brand mark, or logo as
# inline illustrations.
NON_HERO_IMAGES = {
    "amazon_kdp_hardcover.jpg",
    "blue-arrowhead.svg",
    "finnoy_arrowhead_2026_mini-blue.png",
    "finnoybu_arrow_2026_mini-blue.webp",
    "finnoybu-press-logo.png",
}

PRINT_IMAGE_MAX_WIDTH = 1200
PRINT_IMAGE_QUALITY = 85


def prepare_print_images():
    """Convert chapter heroes from public/images/ to grayscale JPEG at
    print-friendly dimensions. The source WebP should already be B&W, but
    `.convert('L')` is a cheap safety net against any color leakage."""
    os.makedirs(PRINT_IMAGES_DIR, exist_ok=True)
    count = 0
    for fname in sorted(os.listdir(PUBLIC_IMAGES)):
        if fname in NON_HERO_IMAGES:
            continue
        ext = os.path.splitext(fname)[1].lower()
        if ext not in (".webp", ".png", ".jpg", ".jpeg"):
            continue
        src = os.path.join(PUBLIC_IMAGES, fname)
        dst = os.path.join(PRINT_IMAGES_DIR, os.path.splitext(fname)[0] + ".jpg")
        if os.path.exists(dst) and os.path.getmtime(dst) > os.path.getmtime(src):
            count += 1
            continue
        img = Image.open(src).convert("L")
        if img.width > PRINT_IMAGE_MAX_WIDTH:
            ratio = PRINT_IMAGE_MAX_WIDTH / img.width
            img = img.resize(
                (PRINT_IMAGE_MAX_WIDTH, int(img.height * ratio)),
                Image.LANCZOS,
            )
        img.save(dst, "JPEG", quality=PRINT_IMAGE_QUALITY, optimize=True)
        count += 1
    return count


def load_chapters():
    chapters = []
    for fname in sorted(os.listdir(CONTENT_DIR)):
        if not fname.endswith(".md"):
            continue
        fpath = os.path.join(CONTENT_DIR, fname)
        with open(fpath, "r", encoding="utf-8") as f:
            text = f.read()

        m = re.match(r"^---\n(.*?)\n---\s*", text, re.DOTALL)
        if not m:
            continue
        meta = yaml.safe_load(m.group(1))
        body = text[m.end():]

        hero = meta.get("hero", {})
        hero_image = hero.get("image", "") if isinstance(hero, dict) else ""

        chapters.append({
            "id": meta.get("id", 999),
            "title": meta.get("title", fname.replace(".md", "").replace("-", " ").title()),
            "body": body,
            "fname": fname,
            "hero_image": hero_image,
        })

    chapters.sort(key=lambda x: x["id"])
    return chapters


def resolve_hero_path(hero_image):
    """Map frontmatter hero.image path (e.g. /images/foo.webp) to the
    grayscale JPEG produced in output/images-print/."""
    if not hero_image:
        return None
    base = os.path.splitext(os.path.basename(hero_image))[0]
    p = os.path.join(PRINT_IMAGES_DIR, base + ".jpg")
    return p if os.path.exists(p) else None


def strip_inline_images(text):
    """Remove inline body image markdown — chapter heroes are inserted
    by build_combined_markdown(), and we don't want duplicates."""
    text = re.sub(r"!\[[^\]]*\]\([^)]+\)\s*\n?", "", text)
    text = re.sub(r"<img[^>]*/?>", "", text)
    return text


def build_combined_markdown(chapters, front_matter_md=""):
    parts = []
    if front_matter_md.strip():
        parts.append(front_matter_md.strip())

    for ch in chapters:
        if ch["id"] == 0:
            heading = f"# {ch['title']}"
        else:
            heading = f"# Chapter {ch['id']}: {ch['title']}"

        body = strip_inline_images(ch["body"].strip())

        hero_md = ""
        hero_path = resolve_hero_path(ch["hero_image"])
        if hero_path:
            # `width=75%` of page width ≈ 35% of the source illustration's
            # natural display size (heroes are ~9.5 in wide at Word's
            # default DPI; full-bleed = 4.5 in printable = ~47% of natural).
            hero_md = f"\n\n![{ch['title']}]({hero_path}){{ width=75% }}\n"

        parts.append(f"{heading}{hero_md}\n\n{body}")
    return "\n\n\n".join(parts)


# --- Post-processing helpers (ported from finnoybu-fiction) ---

def _make_text_p(text, size_pt, bold=False, font="Georgia"):
    bold_xml = "<w:b/>" if bold else ""
    return parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr><w:jc w:val="center"/>'
        f'    <w:rPr><w:rFonts w:ascii="{font}" w:hAnsi="{font}"/>'
        f'      <w:sz w:val="{size_pt * 2}"/>{bold_xml}</w:rPr>'
        f'  </w:pPr>'
        f'  <w:r>'
        f'    <w:rPr><w:rFonts w:ascii="{font}" w:hAnsi="{font}"/>'
        f'      <w:sz w:val="{size_pt * 2}"/>{bold_xml}</w:rPr>'
        f'    <w:t xml:space="preserve">{text}</w:t>'
        f'  </w:r>'
        f'</w:p>'
    )


def _make_blank_p():
    return parse_xml(f'<w:p {nsdecls("w")}/>')


def _make_page_break_p():
    return parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:r><w:br w:type="page"/></w:r>'
        f'</w:p>'
    )


def _set_h2_double_underline(h2_style):
    pPr = h2_style.element.get_or_add_pPr()
    for old in pPr.findall(qn("w:pBdr")):
        pPr.remove(old)
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="double" w:sz="6" w:space="4" w:color="auto"/>'
        f'</w:pBdr>'
    ))


def style_body_and_headings(doc):
    """Body Georgia 11, H1 Calibri 26 bold, H2 Calibri 16 bold + double
    rule, H3 Calibri 12. Page-break-before on H1 and H2 so chapters and
    front-matter sections each open on a new page."""
    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15

    for style_name in ("Body Text", "First Paragraph"):
        try:
            s = doc.styles[style_name]
            s.paragraph_format.line_spacing = 1.15
        except KeyError:
            pass

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(26)
    h1.font.bold = True
    h1.paragraph_format.page_break_before = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.paragraph_format.page_break_before = True
    _set_h2_double_underline(h2)

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)


def set_trade_paperback_page_size(doc):
    """6 x 9 inch trim with 0.75 inch margins (KDP / IngramSpark default)."""
    for section in doc.sections:
        section.page_width = Inches(6.0)
        section.page_height = Inches(9.0)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)


def add_page_numbers_no_first(doc):
    """Centered PAGE field in every section's footer; suppress the number
    on the title-page section's first page only."""
    for i, section in enumerate(doc.sections):
        footer = section.footer
        for p in list(footer.paragraphs):
            p._element.getparent().remove(p._element)
        p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fld_xml = (
            f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE   \\* MERGEFORMAT ">'
            f'  <w:r><w:t>1</w:t></w:r>'
            f'</w:fldSimple>'
        )
        run._element.addnext(parse_xml(fld_xml))

        if i == 0:
            section.different_first_page_header_footer = True
            first_footer = section.first_page_footer
            for p in list(first_footer.paragraphs):
                p._element.getparent().remove(p._element)
            first_footer.add_paragraph()
        else:
            section.different_first_page_header_footer = False


def style_title_page(docx_path):
    """Replace Pandoc's auto-generated title block with a custom title
    page (book title / author / year / logo / publisher), then apply
    body/heading typography, page size, and page numbers."""
    doc = Document(docx_path)
    body = doc.element.body

    pandoc_title_styles = {"Title", "Author", "Subtitle", "Date"}
    to_remove = []
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            break
        if para.style.name in pandoc_title_styles or para.text.strip() == "":
            to_remove.append(para)
        else:
            break
    for p in to_remove:
        p._element.getparent().remove(p._element)

    first_remaining = None
    for para in doc.paragraphs:
        first_remaining = para
        break
    if first_remaining is None:
        print("  WARNING: no anchor paragraph found; skipping title page styling")
        return
    anchor = first_remaining._element

    elements = []
    elements.append(_make_text_p(BOOK_TITLE, size_pt=22, bold=True))
    for _ in range(4):
        elements.append(_make_blank_p())
    elements.append(_make_text_p(AUTHOR, size_pt=14))
    elements.append(_make_blank_p())
    elements.append(_make_text_p(YEAR, size_pt=14))
    elements.append(_make_blank_p())
    elements.append(_make_text_p(TRANSLATOR_LINE, size_pt=11))
    elements.append(_make_text_p(TRANSCRIBER_LINE, size_pt=11))
    for _ in range(4):
        elements.append(_make_blank_p())

    logo_marker_idx = len(elements)
    elements.append(None)
    elements.append(_make_text_p(PUBLISHER, size_pt=14))
    elements.append(_make_page_break_p())

    if os.path.exists(LOGO_PATH):
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = logo_para.add_run()
        run.add_picture(LOGO_PATH, height=Inches(0.75))
        logo_el = logo_para._element
        body.remove(logo_el)
        elements[logo_marker_idx] = logo_el
    else:
        elements[logo_marker_idx] = _make_blank_p()
        print(f"  WARNING: logo not found at {LOGO_PATH}; rendered without it")

    for el in elements:
        anchor.addprevious(el)

    style_body_and_headings(doc)
    set_trade_paperback_page_size(doc)
    add_page_numbers_no_first(doc)

    doc.save(docx_path)
    print(f"  Title page styled: {BOOK_TITLE} / {AUTHOR} / {YEAR} / {PUBLISHER}")
    print(f"  Body: Georgia 11; H1: Calibri 26 bold; H2: Calibri 16 bold + double rule; H3: Calibri 12")
    print(f"  Page-break-before on Heading 1 and Heading 2")
    print(f"  Page size: 6 x 9 in with 0.75 in margins applied to {len(doc.sections)} sections")
    print(f"  Page numbers added (centered footer, suppressed on title page)")


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    print("Preparing print images (grayscale JPEG, 1200px max)...")
    n_imgs = prepare_print_images()
    print(f"  {n_imgs} images ready in {PRINT_IMAGES_DIR}")

    front_matter_md = ""
    if os.path.exists(FRONT_MATTER_PATH):
        with open(FRONT_MATTER_PATH, "r", encoding="utf-8") as f:
            front_matter_md = f.read()
        print(f"  front-matter: {FRONT_MATTER_PATH}")
    else:
        print(f"  (no front-matter at {FRONT_MATTER_PATH}; proceeding without)")

    print(f"Loading chapters from {CONTENT_DIR}...")
    chapters = load_chapters()
    print(f"  Found {len(chapters)} chapters")

    print("Assembling markdown...")
    combined = build_combined_markdown(chapters, front_matter_md)
    combined_md_path = os.path.join(OUTPUT_DIR, "combined.md")
    with open(combined_md_path, "w", encoding="utf-8") as f:
        f.write(combined)
    print(f"  Written to {combined_md_path}")

    output_path = os.path.join(OUTPUT_DIR, OUTPUT_FILE)

    # No --toc: TOC is inserted manually in Word at the position the user
    # controls (matches finnoybu-fiction's print flow). Pandoc's auto-TOC
    # would land at the very top of the document, before the title page.
    # `markdown-implicit_figures`: suppress Pandoc's auto-caption behaviour
    # where an image alone in a paragraph gets its alt text rendered as a
    # caption beneath the figure. Heroes inherit the chapter title as alt
    # text for accessibility, but that produces a duplicated chapter
    # heading in the printed flow.
    pandoc_cmd = [
        "pandoc",
        combined_md_path,
        "-o", output_path,
        "--from", "markdown-implicit_figures",
        "--to", "docx",
        "--standalone",
        f"--metadata=title:{BOOK_TITLE}",
        f"--metadata=author:{AUTHOR}",
    ]

    print("Running Pandoc...")
    result = subprocess.run(pandoc_cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"  Pandoc STDERR: {result.stderr[:500]}")

    if not os.path.exists(output_path):
        print("  ERROR: DOCX not created")
        return

    size_kb = os.path.getsize(output_path) / 1024
    print(f"  Pandoc SUCCESS: {output_path} ({size_kb:.0f} KB)")

    print("Styling title page + post-processing...")
    style_title_page(output_path)
    size_kb = os.path.getsize(output_path) / 1024
    print(f"  Final: {output_path} ({size_kb:.0f} KB)")


if __name__ == "__main__":
    main()
